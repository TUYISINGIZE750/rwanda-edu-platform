from docx import Document
import re

doc = Document('LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC__2024-2025.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Extract all text
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

# Print first 200 lines
for i, line in enumerate(all_text[:200]):
    print(f"{i+1}: {line}")

# Try to find patterns
print("\n\n=== Looking for school patterns ===")
for line in all_text[:100]:
    if any(keyword in line.lower() for keyword in ['tvet', 'tss', 'vtc', 'iprc', 'ecole', 'centre', 'school']):
        print(line)
